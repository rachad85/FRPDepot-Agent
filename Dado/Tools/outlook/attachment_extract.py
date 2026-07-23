#!/usr/bin/env python
"""Extract searchable text from sanitized FRP Depot Outlook attachments.

The script reads only local audit copies. It never writes to Outlook. Content
matching the different-company hard wall is excluded and its local copy purged.
"""

from __future__ import annotations

from collections import Counter, defaultdict
from datetime import datetime, timezone
from email import policy
from email.parser import BytesParser
import hashlib
from html import unescape
import json
from pathlib import Path
import re
import sys
from typing import Any
import zipfile

ROOT = Path(r"C:\FRPDepot")
VENDOR = ROOT / "Dado" / "Tools" / "vendor"
if VENDOR.exists():
    sys.path.append(str(VENDOR))

import fitz
from docx import Document
import openpyxl
import xlrd
from PIL import Image

AUDIT_ROOT = ROOT / "Dado" / "20_Working" / "outlook_mailbox_audit"
FORBIDDEN_NEEDLES = ("troydualam", "troy dualam")
MAX_TEXT_CHARS_PER_FILE = 2_000_000
MAX_ZIP_UNCOMPRESSED = 100 * 1024 * 1024

_ocr_engine = None


class ExtractError(RuntimeError):
    pass


def get_ocr_engine():
    global _ocr_engine
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
    return _ocr_engine


def ocr_image(source: Any) -> str:
    result, _ = get_ocr_engine()(source)
    if not result:
        return ""
    return "\n".join(str(item[1]) for item in result if len(item) > 1).strip()


def forbidden_text(text: str) -> bool:
    lowered = text.casefold()
    return any(needle in lowered for needle in FORBIDDEN_NEEDLES)


def html_to_text(value: str) -> str:
    value = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", value)
    value = re.sub(r"(?i)<br\s*/?>", "\n", value)
    value = re.sub(r"(?i)</p\s*>", "\n", value)
    value = re.sub(r"(?s)<[^>]+>", " ", value)
    value = unescape(value)
    return re.sub(r"[ \t]+", " ", value).strip()


def extract_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    document = fitz.open(path)
    parts: list[str] = []
    ocr_pages = 0
    try:
        for page_number, page in enumerate(document):
            text = page.get_text("text").strip()
            if len(text) < 20:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                text = ocr_image(pixmap.tobytes("png"))
                if text:
                    ocr_pages += 1
            parts.append(f"\n--- Page {page_number + 1} ---\n{text}")
    finally:
        document.close()
    return "".join(parts).strip(), {"pages": len(parts), "ocr_pages": ocr_pages}


def extract_docx(path: Path) -> tuple[str, dict[str, Any]]:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_count = 0
    for table in document.tables:
        table_count += 1
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts), {"paragraphs": len(document.paragraphs), "tables": table_count}


def extract_xlsx(path: Path) -> tuple[str, dict[str, Any]]:
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=False)
    parts: list[str] = []
    row_count = 0
    try:
        for sheet in workbook.worksheets:
            parts.append(f"\n--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    parts.append("\t".join(values))
                    row_count += 1
    finally:
        workbook.close()
    return "\n".join(parts).strip(), {"sheets": len(workbook.sheetnames), "rows": row_count}


def extract_xls(path: Path) -> tuple[str, dict[str, Any]]:
    workbook = xlrd.open_workbook(path, on_demand=True)
    parts: list[str] = []
    row_count = 0
    try:
        for sheet in workbook.sheets():
            parts.append(f"\n--- Sheet: {sheet.name} ---")
            for row_index in range(sheet.nrows):
                values = [str(sheet.cell_value(row_index, column)) for column in range(sheet.ncols)]
                if any(value.strip() for value in values):
                    parts.append("\t".join(values))
                    row_count += 1
    finally:
        workbook.release_resources()
    return "\n".join(parts).strip(), {"sheets": workbook.nsheets, "rows": row_count}


def extract_eml(path: Path) -> tuple[str, dict[str, Any]]:
    message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    parts = [
        f"Subject: {message.get('subject', '')}",
        f"From: {message.get('from', '')}",
        f"To: {message.get('to', '')}",
        f"Date: {message.get('date', '')}",
    ]
    attached_names: list[str] = []
    for part in message.walk():
        if part.is_multipart():
            continue
        filename = part.get_filename()
        if filename:
            attached_names.append(filename)
            continue
        content_type = part.get_content_type()
        if content_type not in {"text/plain", "text/html"}:
            continue
        try:
            content = part.get_content()
        except Exception:
            payload = part.get_payload(decode=True) or b""
            content = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
        parts.append(html_to_text(content) if content_type == "text/html" else str(content))
    return "\n".join(parts), {"embedded_attachments": attached_names}


def extract_zip(path: Path) -> tuple[str, dict[str, Any]]:
    with zipfile.ZipFile(path) as archive:
        infos = archive.infolist()
        total = sum(info.file_size for info in infos)
        if total > MAX_ZIP_UNCOMPRESSED:
            raise ExtractError("ZIP uncompressed size exceeds 100 MB safety limit")
        names = [info.filename for info in infos]
        parts = ["ZIP contents:", *names]
        for info in infos:
            if info.is_dir() or info.file_size > 2 * 1024 * 1024:
                continue
            suffix = Path(info.filename).suffix.casefold()
            if suffix in {".txt", ".csv", ".json", ".xml", ".md"}:
                raw = archive.read(info)
                parts.append(f"\n--- {info.filename} ---\n" + raw.decode("utf-8", errors="replace"))
        return "\n".join(parts), {"entries": len(infos), "uncompressed_bytes": total}


def extract_one(path: Path, content_type: str) -> tuple[str, str, dict[str, Any]]:
    suffix = path.suffix.casefold()
    if content_type == "application/pdf" or suffix == ".pdf":
        text, metadata = extract_pdf(path)
        return text, "pdf_text_and_ocr", metadata
    if suffix == ".docx":
        text, metadata = extract_docx(path)
        return text, "docx", metadata
    if suffix == ".xlsx":
        text, metadata = extract_xlsx(path)
        return text, "xlsx", metadata
    if suffix == ".xls":
        text, metadata = extract_xls(path)
        return text, "xls", metadata
    if content_type == "message/rfc822" or suffix == ".eml" or not suffix:
        text, metadata = extract_eml(path)
        return text, "eml", metadata
    if suffix in {".jpg", ".jpeg", ".png", ".webp", ".tif", ".tiff"}:
        with Image.open(path) as image:
            metadata = {"width": image.width, "height": image.height, "format": image.format}
        return ocr_image(str(path)), "image_ocr", metadata
    if suffix == ".zip":
        text, metadata = extract_zip(path)
        return text, "zip_inventory", metadata
    if suffix in {".step", ".stp", ".ics", ".txt", ".csv", ".json", ".xml"}:
        return path.read_text(encoding="utf-8", errors="replace"), "plain_text", {}
    if suffix == ".mp3":
        return "", "audio_pending_transcription", {"bytes": path.stat().st_size}
    return "", "metadata_only", {"bytes": path.stat().st_size}


def extract_run(run_dir: Path) -> Path:
    index_path = run_dir / "attachments.jsonl"
    rows = [
        json.loads(line)
        for line in index_path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    downloaded = [row for row in rows if row.get("status") == "downloaded" and row.get("local_path")]
    by_hash: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in downloaded:
        path = Path(str(row["local_path"]))
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        row["sha256"] = digest
        by_hash[digest].append(row)

    output_path = run_dir / "attachment_text.jsonl"
    results: list[dict[str, Any]] = []
    excluded_hashes: set[str] = set()
    methods = Counter()
    errors = 0

    for position, (digest, linked_rows) in enumerate(by_hash.items(), start=1):
        representative = linked_rows[0]
        path = Path(str(representative["local_path"]))
        content_type = str(representative.get("content_type") or "")
        name = str(representative.get("name") or path.name)
        result: dict[str, Any] = {
            "sha256": digest,
            "name": name,
            "content_type": content_type,
            "message_ids": sorted({str(row.get("message_id") or "") for row in linked_rows}),
            "copies": len(linked_rows),
            "method": None,
            "metadata": {},
            "text": "",
            "error": None,
        }
        try:
            raw_probe = path.read_bytes().lower()
            raw_blocked = any(needle.encode("utf-8") in raw_probe for needle in FORBIDDEN_NEEDLES)
            name_blocked = forbidden_text(name)
            if raw_blocked or name_blocked:
                excluded_hashes.add(digest)
                result["method"] = "excluded_company_wall"
            else:
                text, method, metadata = extract_one(path, content_type)
                if forbidden_text(text):
                    excluded_hashes.add(digest)
                    result["method"] = "excluded_company_wall"
                else:
                    result["method"] = method
                    result["metadata"] = metadata
                    result["text"] = text[:MAX_TEXT_CHARS_PER_FILE]
        except Exception as exc:
            result["method"] = "error"
            result["error"] = f"{type(exc).__name__}: {exc}"
            errors += 1
        methods[result["method"]] += 1
        results.append(result)
        if position % 25 == 0:
            print(f"Unique attachments processed: {position}/{len(by_hash)}", flush=True)

    if excluded_hashes:
        retained_rows: list[dict[str, Any]] = []
        for row in rows:
            if row.get("sha256") in excluded_hashes:
                local_path = row.get("local_path")
                if local_path:
                    path = Path(str(local_path))
                    if path.exists():
                        path.unlink()
            else:
                retained_rows.append(row)
        with index_path.open("w", encoding="utf-8", newline="\n") as handle:
            for row in retained_rows:
                handle.write(json.dumps(row, ensure_ascii=False) + "\n")
        results = [result for result in results if result["sha256"] not in excluded_hashes]

    with output_path.open("w", encoding="utf-8", newline="\n") as handle:
        for result in results:
            handle.write(json.dumps(result, ensure_ascii=False) + "\n")

    summary = {
        "completed_utc": datetime.now(timezone.utc).isoformat(),
        "downloaded_attachment_rows": len(downloaded),
        "unique_attachment_hashes": len(by_hash),
        "unique_attachments_retained": len(results),
        "unique_attachments_excluded_company_wall": len(excluded_hashes),
        "methods": dict(sorted(methods.items())),
        "errors": errors,
        "output": str(output_path),
    }
    summary_path = run_dir / "attachment_text_summary.json"
    summary_path.write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8")
    outlook_summary_path = run_dir / "summary.json"
    outlook_summary = json.loads(outlook_summary_path.read_text(encoding="utf-8"))
    outlook_summary["attachment_text_summary"] = str(summary_path)
    outlook_summary["attachments_excluded_company_wall_after_content_scan"] = len(excluded_hashes)
    outlook_summary_path.write_text(json.dumps(outlook_summary, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(summary, indent=2))
    return output_path


def main() -> int:
    latest = (AUDIT_ROOT / "LATEST.txt").read_text(encoding="utf-8").strip()
    run_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(latest)
    try:
        extract_run(run_dir)
        return 0
    except (OSError, ExtractError, json.JSONDecodeError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
